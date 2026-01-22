from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ========== PAGE DE GARDE ==========
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RAPPORT E4")
run.bold = True
run.font.size = Pt(28)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Realisation d'un Service Numerique")
run.bold = True
run.font.size = Pt(20)

doc.add_paragraph()
doc.add_paragraph()

project = doc.add_paragraph()
project.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = project.add_run("Projet WTTJ")
run.font.size = Pt(18)
run.bold = True

project_desc = doc.add_paragraph()
project_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = project_desc.add_run("Plateforme de Collecte et Mise a Disposition\ndes Offres d'Emploi Data")
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Candidat(e): ").bold = True
info.add_run("[Votre Nom]\n")
info.add_run("Formation: ").bold = True
info.add_run("Expert en Donnees Massives (RNCP 37638)\n")
info.add_run("Date: ").bold = True
info.add_run("Janvier 2026")

doc.add_page_break()

# ========== TABLE DES MATIERES ==========
doc.add_heading("Table des matieres", level=1)
toc_items = [
    "1. Presentation du projet et contexte",
    "2. Specifications techniques",
    "3. Extraction automatisee des donnees (C8)",
    "4. Requetes SQL d'extraction (C9)",
    "5. Agregation et nettoyage des donnees (C10)",
    "6. Creation de la base de donnees (C11)",
    "7. API REST et mise a disposition (C12)",
    "8. Automatisation du pipeline",
    "9. Conformite RGPD et vigilances",
    "10. Conclusion et perspectives",
    "Annexes"
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ========== SECTION 1 ==========
doc.add_heading("1. Presentation du projet et contexte", level=1)

doc.add_heading("1.1 Contexte organisationnel", level=2)
doc.add_paragraph(
    "Le projet WTTJ (Welcome To The Jungle) s'inscrit dans le cadre d'une mission de veille "
    "et d'analyse du marche de l'emploi dans le secteur de la Data en France. L'objectif est "
    "de fournir aux equipes RH et aux analystes metier un acces structure aux offres d'emploi "
    "du domaine Data (Data Analyst, Data Engineer, Data Scientist, etc.)."
)

doc.add_heading("1.2 Expression du besoin", level=2)
doc.add_paragraph(
    "Problematique: Les donnees d'offres d'emploi sont dispersees sur differentes plateformes "
    "et ne sont pas exploitables directement pour des analyses statistiques ou du reporting."
)

p = doc.add_paragraph()
p.add_run("Objectifs fonctionnels:").bold = True
doc.add_paragraph("Collecter automatiquement les offres d'emploi Data depuis Welcome To The Jungle", style='List Bullet')
doc.add_paragraph("Structurer et nettoyer les donnees collectees", style='List Bullet')
doc.add_paragraph("Stocker les donnees dans une base relationnelle", style='List Bullet')
doc.add_paragraph("Exposer les donnees via une API REST securisee", style='List Bullet')

p = doc.add_paragraph()
p.add_run("Objectifs techniques:").bold = True
doc.add_paragraph("Automatiser l'extraction via scraping et appels API", style='List Bullet')
doc.add_paragraph("Implementer une base de donnees relationnelle sur Azure SQL Server", style='List Bullet')
doc.add_paragraph("Developper une API REST avec authentification JWT", style='List Bullet')
doc.add_paragraph("Assurer la conformite RGPD", style='List Bullet')

doc.add_heading("1.3 Environnement technique", level=2)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Composant"
table.rows[0].cells[1].text = "Technologie"
data = [
    ("Scraping", "Selenium, BeautifulSoup"),
    ("Langage", "Python 3.12"),
    ("Base de donnees", "Azure SQL Server"),
    ("ORM", "SQLAlchemy"),
    ("API", "FastAPI, Uvicorn"),
    ("Authentification", "JWT (python-jose)"),
    ("Versioning", "Git")
]
for i, (comp, tech) in enumerate(data):
    table.rows[i+1].cells[0].text = comp
    table.rows[i+1].cells[1].text = tech

doc.add_page_break()

# ========== SECTION 2 ==========
doc.add_heading("2. Specifications techniques", level=1)

doc.add_heading("2.1 Architecture globale", level=2)
doc.add_paragraph("Le systeme est compose de trois modules principaux:")
doc.add_paragraph("[Sources Web: WTTJ.com] --> [Traitement: Scraper Python] --> [Stockage: Azure SQL] --> [Exposition: API REST]")

doc.add_heading("2.2 Sources de donnees", level=2)
doc.add_paragraph("L'extraction combine deux sources:")
doc.add_paragraph("Scraping web (Selenium): Collecte des liens d'offres depuis les pages de recherche", style='List Bullet')
doc.add_paragraph("API WTTJ: Enrichissement des donnees via l'API interne de Welcome To The Jungle", style='List Bullet')

doc.add_heading("2.3 Volumetrie", level=2)
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Metrique"
table.rows[0].cells[1].text = "Valeur"
vol_data = [
    ("Offres collectees", "3 695 jobs uniques"),
    ("Entreprises", "1 200+"),
    ("Localisations", "500+"),
    ("Skills", "6 943"),
    ("Tools", "9 307"),
    ("Benefits", "54 375")
]
for i, (m, v) in enumerate(vol_data):
    table.rows[i+1].cells[0].text = m
    table.rows[i+1].cells[1].text = v

doc.add_page_break()

# ========== SECTION 3 (C8) ==========
doc.add_heading("3. Extraction automatisee des donnees (C8)", level=1)

doc.add_heading("3.1 Script de scraping (job_scraper.py)", level=2)
doc.add_paragraph("Le script utilise Selenium en mode headless pour parcourir les pages de recherche WTTJ. Il collecte les liens des offres d'emploi en parcourant toutes les pages de resultats.")

p = doc.add_paragraph()
p.add_run("Fonctionnement du script:").bold = True
doc.add_paragraph("Configuration Chrome en mode headless avec rotation des User-Agents", style='List Bullet')
doc.add_paragraph("Parcours pagine des resultats de recherche", style='List Bullet')
doc.add_paragraph("Extraction des liens via selecteurs CSS", style='List Bullet')
doc.add_paragraph("Stockage dans un DataFrame pandas", style='List Bullet')

doc.add_heading("3.2 Enrichissement via API (api_scraper.py)", level=2)
doc.add_paragraph("Chaque lien collecte est enrichi via l'API interne de WTTJ pour recuperer: informations detaillees du poste, skills et tools requis, benefits proposes, informations sur l'entreprise.")

doc.add_heading("3.3 Configuration", level=2)
doc.add_paragraph("Mots-cles configures: data analyst, data engineer, data scientist, machine learning, business intelligence, data architect, intelligence artificielle")
doc.add_paragraph("Zone geographique: France (FR)")

doc.add_page_break()

# ========== SECTION 4 (C9) ==========
doc.add_heading("4. Requetes SQL d'extraction (C9)", level=1)

doc.add_heading("4.1 Requetes avec SQLAlchemy ORM", level=2)
doc.add_paragraph("Les requetes utilisent l'ORM SQLAlchemy avec joinedload pour optimiser les performances et eviter le probleme N+1.")

p = doc.add_paragraph()
p.add_run("Exemple SQLAlchemy - Extraction des jobs avec relations:").bold = True
code1 = doc.add_paragraph()
code1.add_run("""jobs = (
    db.query(Job)
    .options(
        joinedload(Job.company),
        joinedload(Job.location),
        joinedload(Job.skills),
        joinedload(Job.tools),
        joinedload(Job.benefits)
    )
    .order_by(Job.job_reference)
    .offset(skip).limit(limit)
    .all()
)""").font.size = Pt(9)

doc.add_heading("4.2 Exemples de requetes SQL natives", level=2)

p = doc.add_paragraph()
p.add_run("Requete 1 - Extraction des offres avec entreprise et localisation:").bold = True
sql1 = doc.add_paragraph()
sql1.add_run("""SELECT j.job_reference, j.poste, j.salary_min, j.salary_max,
       c.name AS company_name, c.industry,
       l.city, l.country_code
FROM jobs j
LEFT JOIN companies c ON j.company_id = c.id
LEFT JOIN locations l ON j.location_id = l.id
ORDER BY j.published_at DESC
OFFSET 0 ROWS FETCH NEXT 100 ROWS ONLY;""").font.size = Pt(9)

p = doc.add_paragraph()
p.add_run("Requete 2 - Agregation des skills par offre d'emploi:").bold = True
sql2 = doc.add_paragraph()
sql2.add_run("""SELECT j.job_reference, j.poste,
       STRING_AGG(s.skill, ', ') AS competences
FROM jobs j
INNER JOIN skills s ON j.job_reference = s.job_reference
GROUP BY j.job_reference, j.poste
ORDER BY j.job_reference;""").font.size = Pt(9)

p = doc.add_paragraph()
p.add_run("Requete 3 - Statistiques salariales par type de contrat:").bold = True
sql3 = doc.add_paragraph()
sql3.add_run("""SELECT contract_type,
       COUNT(*) AS nb_offres,
       AVG(salary_min) AS salaire_min_moyen,
       AVG(salary_max) AS salaire_max_moyen,
       MIN(salary_min) AS salaire_min,
       MAX(salary_max) AS salaire_max
FROM jobs
WHERE salary_min IS NOT NULL AND salary_max IS NOT NULL
GROUP BY contract_type
ORDER BY nb_offres DESC;""").font.size = Pt(9)

p = doc.add_paragraph()
p.add_run("Requete 4 - Top 10 des skills les plus demandes:").bold = True
sql4 = doc.add_paragraph()
sql4.add_run("""SELECT TOP 10 skill, COUNT(*) AS occurrences
FROM skills
GROUP BY skill
ORDER BY occurrences DESC;""").font.size = Pt(9)

p = doc.add_paragraph()
p.add_run("Requete 5 - Offres par ville avec nombre de skills requis:").bold = True
sql5 = doc.add_paragraph()
sql5.add_run("""SELECT l.city, COUNT(DISTINCT j.job_reference) AS nb_offres,
       COUNT(s.id) AS total_skills,
       AVG(j.salary_min) AS salaire_moyen
FROM jobs j
INNER JOIN locations l ON j.location_id = l.id
LEFT JOIN skills s ON j.job_reference = s.job_reference
GROUP BY l.city
HAVING COUNT(DISTINCT j.job_reference) > 5
ORDER BY nb_offres DESC;""").font.size = Pt(9)

p = doc.add_paragraph()
p.add_run("Requete 6 - Recherche full-text sur les profils:").bold = True
sql6 = doc.add_paragraph()
sql6.add_run("""SELECT job_reference, poste, company_id
FROM jobs
WHERE profile LIKE '%Python%'
   OR profile LIKE '%SQL%'
ORDER BY published_at DESC;""").font.size = Pt(9)

doc.add_heading("4.3 Optimisations mises en place", level=2)
doc.add_paragraph("Utilisation de joinedload pour le chargement eager des relations (evite N+1)", style='List Bullet')
doc.add_paragraph("Index sur les cles primaires (job_reference) et etrangeres (company_id, location_id)", style='List Bullet')
doc.add_paragraph("Pagination avec OFFSET/FETCH pour SQL Server (obligatoire avec ORDER BY)", style='List Bullet')
doc.add_paragraph("Limitation des resultats avec FETCH NEXT pour eviter les timeouts", style='List Bullet')
doc.add_paragraph("Pool de connexions SQLAlchemy avec pool_pre_ping pour la resilience", style='List Bullet')

doc.add_page_break()

# ========== SECTION 5 (C10) ==========
doc.add_heading("5. Agregation et nettoyage des donnees (C10)", level=1)

doc.add_heading("5.1 Pipeline d'agregation", level=2)
doc.add_paragraph("Le script insert_clean_data.py realise l'agregation et le nettoyage des donnees collectees avant insertion en base de donnees.")

doc.add_heading("5.2 Regles de nettoyage", level=2)
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Champ"
table.rows[0].cells[1].text = "Traitement"
clean_data = [
    ("skills/tools", "Split par virgule, trim des espaces"),
    ("salary_min/max", "Conversion en entier, gestion des nulls"),
    ("published_at", "Parsing format ISO 8601"),
    ("cover_letter", "Normalisation en booleen"),
    ("Doublons", "Deduplication sur job_reference")
]
for i, (c, t) in enumerate(clean_data):
    table.rows[i+1].cells[0].text = c
    table.rows[i+1].cells[1].text = t

doc.add_heading("5.3 Gestion des doublons", level=2)
doc.add_paragraph("Avant chaque insertion, verification de l'existence du job_reference en base. Les enregistrements existants sont ignores pour eviter les doublons.")

doc.add_page_break()

# ========== SECTION 6 (C11) ==========
doc.add_heading("6. Creation de la base de donnees (C11)", level=1)

doc.add_heading("6.1 Modele Conceptuel de Donnees (MCD)", level=2)
doc.add_paragraph("Entites et relations principales:")
doc.add_paragraph("COMPANY (1,n) --- emploie --- (0,n) JOB", style='List Bullet')
doc.add_paragraph("LOCATION (1,n) --- localise --- (0,n) JOB", style='List Bullet')
doc.add_paragraph("JOB (1,1) --- possede --- (0,n) SKILL", style='List Bullet')
doc.add_paragraph("JOB (1,1) --- utilise --- (0,n) TOOL", style='List Bullet')
doc.add_paragraph("JOB (1,1) --- offre --- (0,n) BENEFIT", style='List Bullet')
doc.add_paragraph("JOB (1,1) --- a --- (0,1) MEDIA", style='List Bullet')

doc.add_heading("6.2 Tables creees", level=2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Table"
table.rows[0].cells[1].text = "Cle primaire"
table.rows[0].cells[2].text = "Description"
tables_data = [
    ("companies", "id (auto)", "Entreprises"),
    ("locations", "id (auto)", "Localisations"),
    ("jobs", "job_reference", "Offres d'emploi"),
    ("skills", "id (auto)", "Competences"),
    ("tools", "id (auto)", "Outils techniques"),
    ("benefits", "id (auto)", "Avantages"),
    ("media", "id (auto)", "Reseaux sociaux")
]
for i, (t, pk, desc) in enumerate(tables_data):
    table.rows[i+1].cells[0].text = t
    table.rows[i+1].cells[1].text = pk
    table.rows[i+1].cells[2].text = desc

doc.add_heading("6.3 Configuration Azure SQL Server", level=2)
doc.add_paragraph("Base hebergee sur Azure SQL Server avec connexion chiffree TLS, timeout de 30 secondes, et pool de connexions.")

doc.add_page_break()

# ========== SECTION 7 (C12) ==========
doc.add_heading("7. API REST et mise a disposition (C12)", level=1)

doc.add_heading("7.1 Architecture de l'API", level=2)
doc.add_paragraph("API construite avec FastAPI exposant les endpoints suivants:")

table = doc.add_table(rows=9, cols=4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Methode"
table.rows[0].cells[1].text = "Endpoint"
table.rows[0].cells[2].text = "Auth"
table.rows[0].cells[3].text = "Description"
api_data = [
    ("GET", "/", "Non", "Message de bienvenue"),
    ("GET", "/health", "Non", "Status de l'API"),
    ("POST", "/auth/login", "Non", "Authentification JWT"),
    ("GET", "/jobs/", "Oui", "Liste des offres"),
    ("GET", "/companies/", "Non", "Liste des entreprises"),
    ("GET", "/locations/", "Non", "Liste des localisations"),
    ("GET", "/skills/", "Non", "Liste des competences"),
    ("GET", "/tools/", "Non", "Liste des outils")
]
for i, (m, e, a, d) in enumerate(api_data):
    table.rows[i+1].cells[0].text = m
    table.rows[i+1].cells[1].text = e
    table.rows[i+1].cells[2].text = a
    table.rows[i+1].cells[3].text = d

doc.add_heading("7.2 Authentification JWT", level=2)
doc.add_paragraph("Tokens JWT avec algorithme HS256 et expiration de 60 minutes. L'endpoint /jobs/ est protege et necessite un token valide.")

doc.add_heading("7.3 Documentation OpenAPI", level=2)
doc.add_paragraph("Documentation Swagger auto-generee accessible sur /docs avec schemas de requetes/reponses et exemples d'appels interactifs.")

doc.add_page_break()

# ========== SECTION 8 (Automatisation) ==========
doc.add_heading("8. Automatisation du pipeline", level=1)

doc.add_heading("8.1 Architecture du pipeline automatise", level=2)
doc.add_paragraph(
    "Un pipeline automatise a ete mis en place pour maintenir les donnees a jour. "
    "Il orchestre l'ensemble du processus: scraping, nettoyage, insertion en base, et archivage des fichiers."
)

p = doc.add_paragraph()
p.add_run("Composants du pipeline:").bold = True
doc.add_paragraph("run_pipeline.py: Script orchestrateur principal", style='List Bullet')
doc.add_paragraph("scheduler.py: Planificateur Python avec la librairie schedule", style='List Bullet')
doc.add_paragraph("Windows Task Scheduler: Alternative native Windows", style='List Bullet')

doc.add_heading("8.2 Etapes du pipeline", level=2)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Etape"
table.rows[0].cells[1].text = "Script"
table.rows[0].cells[2].text = "Description"
pipeline_data = [
    ("1. Scraping", "job_scraper.py + api_scraper.py", "Collecte des offres via Selenium et API WTTJ"),
    ("2. Nettoyage", "data_cleaner.py", "Suppression HTML, doublons, normalisation"),
    ("3. Insertion", "insert_clean_data.py", "Insertion en base avec gestion des doublons"),
    ("4. Archivage", "run_pipeline.py", "Archive CSV et supprime anciens fichiers")
]
for i, (etape, script, desc) in enumerate(pipeline_data):
    table.rows[i+1].cells[0].text = etape
    table.rows[i+1].cells[1].text = script
    table.rows[i+1].cells[2].text = desc

doc.add_heading("8.3 Planification hebdomadaire", level=2)
doc.add_paragraph(
    "Le pipeline est configure pour s'executer automatiquement chaque semaine (dimanche a 02h00) "
    "afin de collecter les nouvelles offres d'emploi tout en evitant les periodes de forte charge."
)

p = doc.add_paragraph()
p.add_run("Options de planification:").bold = True
doc.add_paragraph("Scheduler Python: python scheduler.py (tourne en continu)", style='List Bullet')
doc.add_paragraph("Windows Task Scheduler: Via script PowerShell setup_scheduled_task.ps1", style='List Bullet')
doc.add_paragraph("Execution manuelle: python run_pipeline.py", style='List Bullet')

doc.add_heading("8.4 Gestion des logs et archivage", level=2)
doc.add_paragraph("Logs detailles dans le dossier logs/ avec horodatage", style='List Bullet')
doc.add_paragraph("Archivage automatique des CSV dans data/archive/", style='List Bullet')
doc.add_paragraph("Rotation des archives: conservation des 4 dernieres semaines", style='List Bullet')
doc.add_paragraph("Gestion des erreurs avec rollback en cas d'echec", style='List Bullet')

doc.add_page_break()

# ========== SECTION 9 ==========
doc.add_heading("9. Conformite RGPD et vigilances", level=1)

doc.add_heading("8.1 Analyse des donnees personnelles", level=2)
doc.add_paragraph("Les donnees collectees sont des donnees publiques d'entreprises et d'offres d'emploi. Aucune donnee personnelle de candidats n'est collectee ou stockee.")

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Donnee"
table.rows[0].cells[1].text = "Classification"
table.rows[0].cells[2].text = "Traitement"
rgpd_data = [
    ("Nom entreprise", "Non personnel", "Stockage direct"),
    ("Adresse bureau", "Non personnel", "Stockage direct"),
    ("Coordonnees GPS", "Non personnel", "Stockage direct"),
    ("URL offre", "Non personnel", "Stockage direct")
]
for i, (d, c, t) in enumerate(rgpd_data):
    table.rows[i+1].cells[0].text = d
    table.rows[i+1].cells[1].text = c
    table.rows[i+1].cells[2].text = t

doc.add_heading("8.2 Mesures de securite", level=2)
doc.add_paragraph("Authentification JWT avec expiration 1 heure", style='List Bullet')
doc.add_paragraph("Chiffrement HTTPS/TLS pour l'API", style='List Bullet')
doc.add_paragraph("Connexion chiffree a Azure SQL Server", style='List Bullet')
doc.add_paragraph("Secrets stockes en variables d'environnement (.env non versionne)", style='List Bullet')

doc.add_heading("8.3 Difficultes rencontrees", level=2)
doc.add_paragraph("Blocage anti-scraping: Resolu par rotation des User-Agents", style='List Bullet')
doc.add_paragraph("Timeout connexion Azure: Resolu par configuration timeout 30s", style='List Bullet')
doc.add_paragraph("Doublons dans les donnees: Resolu par deduplication sur cle primaire", style='List Bullet')

doc.add_page_break()

# ========== SECTION 10 ==========
doc.add_heading("10. Conclusion et perspectives", level=1)

doc.add_heading("10.1 Bilan technique", level=2)
doc.add_paragraph("Le projet WTTJ a permis de mettre en oeuvre les competences du bloc 2:")

table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = "Competence"
table.rows[0].cells[1].text = "Realisation"
bilan_data = [
    ("C8 - Extraction automatisee", "Scraping Selenium + API REST"),
    ("C9 - Requetes SQL", "SQLAlchemy ORM + requetes natives"),
    ("C10 - Agregation", "Pipeline pandas + nettoyage"),
    ("C11 - Base de donnees", "Azure SQL Server + modeles MERISE"),
    ("C12 - API REST", "FastAPI + JWT + OpenAPI"),
    ("Automatisation", "Pipeline hebdomadaire + scheduler Python")
]
for i, (c, r) in enumerate(bilan_data):
    table.rows[i+1].cells[0].text = c
    table.rows[i+1].cells[1].text = r

doc.add_heading("10.2 Apprentissages", level=2)
doc.add_paragraph("Maitrise du scraping web avec Selenium", style='List Bullet')
doc.add_paragraph("Conception de bases de donnees relationnelles", style='List Bullet')
doc.add_paragraph("Developpement d'API REST securisees", style='List Bullet')
doc.add_paragraph("Automatisation de pipelines de donnees", style='List Bullet')
doc.add_paragraph("Gestion de projet en methodologie Agile", style='List Bullet')

doc.add_heading("10.3 Ameliorations envisagees", level=2)
doc.add_paragraph("Cache Redis pour ameliorer les performances de l'API", style='List Bullet')
doc.add_paragraph("Dashboard de visualisation des donnees (Power BI / Grafana)", style='List Bullet')
doc.add_paragraph("Modele ML de recommandation d'offres", style='List Bullet')
doc.add_paragraph("Alertes email en cas d'echec du pipeline", style='List Bullet')

doc.add_page_break()

# ========== ANNEXES ==========
doc.add_heading("ANNEXES", level=1)

doc.add_heading("Annexe A - Schema de la base de donnees", level=2)
doc.add_paragraph("Voir document: docs/Annexe_A_Schema_BDD.md")
doc.add_paragraph("Contient: MCD, MLD, MPD et dictionnaire des donnees")

doc.add_heading("Annexe B - Documentation API", level=2)
doc.add_paragraph("Documentation Swagger accessible sur: http://127.0.0.1:8000/docs")
doc.add_paragraph("Voir document: docs/Annexe_B_Documentation_API.md")

doc.add_heading("Annexe C - Depot Git", level=2)
doc.add_paragraph("URL du depot: [A completer avec votre URL GitHub/GitLab]")

doc.add_heading("Annexe D - Captures d'ecran", level=2)
doc.add_paragraph("[Ajouter ici vos captures d'ecran:]")
doc.add_paragraph("- Interface Swagger UI", style='List Bullet')
doc.add_paragraph("- Tests Postman", style='List Bullet')
doc.add_paragraph("- Azure Portal (base de donnees)", style='List Bullet')
doc.add_paragraph("- Execution des scripts", style='List Bullet')

# Sauvegarde
doc.save('docs/Rapport_E4_WTTJ_v3.docx')
print("Rapport Word genere avec succes: docs/Rapport_E4_WTTJ_v3.docx")
